"""
report_generator.py
--------------------
Clean, branded DOCX report generator for the MiffTech IAM/GRC Assistant.
Reproduces the ClearPath assessment design using python-docx (pure Python),
so it drops straight into a Streamlit app.

Usage in Streamlit:
    from report_generator import build_report
    import io

    findings = [
        {
            "id": "F-01",
            "title": "Critical Isolation Failure Reported as Closed",
            "severity": "Critical",
            "observation": "The executive overview reports the risk as closed...",
            "assessment": "The application-layer check is now the sole control...",
            "traceability": [
                "Executive Overview 2.1 - reports risk remediated",
                "Pen Test PT-2025-007 - storage layer still open",
            ],
            "position": "Do not report as closed until the storage layer is fixed.",
        },
        # ...more findings
    ]

    doc_bytes = build_report(
        title="ClearPath Financial Analytics Platform",
        subtitle="Cross-Document Findings, Root Cause, and Remediation Plan",
        prepared_by="Tyrell Mifflin - IAM / GRC Consulting",
        exec_summary="...",
        findings=findings,
        root_cause="...",
    )

    st.download_button(
        "Download Report",
        data=doc_bytes,
        file_name="Security_Assessment.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ---------------------------------------------------------------- palette
NAVY = RGBColor(0x1F, 0x33, 0x52)
ACCENT = RGBColor(0x8A, 0x1C, 0x1C)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
INK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX = "1F3352"
ACCENT_HEX = "8A1C1C"
LIGHT_HEX = "EEF1F5"
ZEBRA_HEX = "F6F8FA"
BORDER_HEX = "D8DDE4"

FONT = "Calibri"

# severity chip colors
SEV_COLORS = {
    "critical": "9B1C1C",
    "high": "B45309",
    "medium": "6B4E14",
    "low": "4A5568",
    "info": "5A5A5A",
}


# ---------------------------------------------------------------- low-level helpers
def _shade(cell, hex_fill):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=90, bottom=90, left=120, right=120):
    """Set internal padding of a cell (in twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tcPr.append(m)


def _cell_border(cell, **kwargs):
    """
    Set borders on a single cell. kwargs like top/bottom/left/right,
    each a dict: {"sz": 24, "color": "1F3352", "val": "single"}.
    Left border thick + others none = the callout-box look.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = kwargs.get(edge)
        node = OxmlElement(f"w:{edge}")
        if spec:
            node.set(qn("w:val"), spec.get("val", "single"))
            node.set(qn("w:sz"), str(spec.get("sz", 4)))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), spec.get("color", "000000"))
        else:
            node.set(qn("w:val"), "nil")
        borders.append(node)
    tcPr.append(borders)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tblPr.append(borders)


def _grid_borders(table, color=BORDER_HEX, sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(sz))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tblPr.append(borders)


def _para_border_bottom(paragraph, color=NAVY_HEX, sz=8):
    """Horizontal rule under a paragraph - used for headings and cover."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(paragraph, text, size=10.5, bold=False, italic=False, color=INK, caps=False):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    if caps:
        r.font.all_caps = True
    return r


# ---------------------------------------------------------------- block builders
def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(8)
    _run(p, text, size=15, bold=True, color=NAVY)
    _para_border_bottom(p, NAVY_HEX, 8)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(5)
    _run(p, text, size=12, bold=True, color=NAVY)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(9)
    pf.space_after = Pt(4)
    _run(p, text, size=11, bold=True, color=ACCENT)
    return p


def body(doc, text, italic=False, color=INK):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    _run(p, text, size=10.5, italic=italic, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    for r in p.runs:
        r.font.name = FONT
        r.font.size = Pt(10.5)
    _run(p, text, size=10.5)
    return p


def callout(doc, label, text, color_hex=ACCENT_HEX):
    """Single-cell shaded box with a thick colored left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, LIGHT_HEX)
    _set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    _cell_border(
        cell,
        left={"sz": 24, "color": color_hex, "val": "single"},
    )
    # label line
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    _run(p0, label.upper(), size=9.5, bold=True, color=RGBColor.from_string(color_hex), caps=True)
    # body line
    p1 = cell.add_paragraph()
    p1.paragraph_format.line_spacing = 1.15
    _run(p1, text, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def data_table(doc, headers, rows, widths=None):
    """Zebra-striped table with a navy header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _grid_borders(table, BORDER_HEX, 4)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], NAVY_HEX)
        _set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        _run(p, h, size=9.5, bold=True, color=WHITE)

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[ci], ZEBRA_HEX)
            _set_cell_margins(cells[ci])
            p = cells[ci].paragraphs[0]
            _run(p, str(val), size=9.5, bold=(ci == 0))

    if widths:
        for ci, w in enumerate(widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    return table


def finding_header(doc, fid, title, severity):
    """Navy title bar with a colored severity chip on the right."""
    sev_hex = SEV_COLORS.get(severity.lower(), "5A5A5A")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(table)

    left = table.cell(0, 0)
    _shade(left, NAVY_HEX)
    _set_cell_margins(left, top=120, bottom=120, left=160, right=120)
    p = left.paragraphs[0]
    _run(p, f"{fid}  ", size=10, bold=True, color=RGBColor(0x9F, 0xB4, 0xD4))
    _run(p, title, size=11, bold=True, color=WHITE)

    right = table.cell(0, 1)
    _shade(right, sev_hex)
    _set_cell_margins(right, top=120, bottom=120, left=100, right=100)
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pr, severity.upper(), size=9.5, bold=True, color=WHITE, caps=True)

    left.width = Inches(5.6)
    right.width = Inches(1.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def _add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Security Assessment   |   Page ", size=8, color=GREY)
    # PAGE field
    run = p.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    run._r.append(fldSimple)
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ---------------------------------------------------------------- main entry point
def build_report(
    title,
    subtitle,
    prepared_by,
    exec_summary,
    findings,
    root_cause,
    classification="Confidential - Restricted",
    bottom_line=None,
):
    """
    Returns DOCX file bytes ready for st.download_button.

    findings: list of dicts with keys:
        id, title, severity, observation, assessment (optional),
        traceability (list), position (optional)
    """
    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)

    # page geometry: US Letter, 0.75in margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    _add_page_number_footer(sec)

    # ---- COVER ----
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    _run(p, "INDEPENDENT SECURITY ASSESSMENT", size=10, bold=True, color=ACCENT, caps=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, title, size=22, bold=True, color=NAVY)
    p = doc.add_paragraph()
    _run(p, subtitle, size=13, color=GREY)
    _para_border_bottom(p, ACCENT_HEX, 12)
    doc.add_paragraph()

    data_table(
        doc,
        ["Field", "Detail"],
        [
            ["Prepared by", prepared_by],
            ["Method", "Trust-layered cross-document reconciliation"],
            ["Basis of finding", "Each finding traceable to named source documents"],
            ["Classification", classification],
        ],
        widths=[1.8, 5.1],
    )
    doc.add_paragraph()
    if bottom_line:
        callout(doc, "Bottom line", bottom_line, ACCENT_HEX)

    doc.add_page_break()

    # ---- EXEC SUMMARY ----
    h1(doc, "1. Executive Summary")
    for para in exec_summary.split("\n\n"):
        if para.strip():
            body(doc, para.strip())

    # findings-at-a-glance table
    if findings:
        doc.add_paragraph()
        h2(doc, "Findings at a Glance")
        rows = [[f["id"], f["title"], f.get("severity", "")] for f in findings]
        data_table(doc, ["#", "Finding", "Rating"], rows, widths=[0.8, 5.0, 1.1])

    doc.add_page_break()

    # ---- DETAILED FINDINGS ----
    h1(doc, "2. Detailed Findings")
    doc.add_paragraph()
    for f in findings:
        finding_header(doc, f["id"], f["title"], f.get("severity", "Info"))
        if f.get("observation"):
            h3(doc, "Observation")
            body(doc, f["observation"])
        if f.get("assessment"):
            h3(doc, "Assessment")
            body(doc, f["assessment"])
        if f.get("traceability"):
            h3(doc, "Traceability")
            for t in f["traceability"]:
                bullet(doc, t)
        if f.get("position"):
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            callout(doc, "Position", f["position"], ACCENT_HEX)
        doc.add_paragraph()

    # ---- ROOT CAUSE ----
    doc.add_page_break()
    h1(doc, "3. Root Cause")
    for para in root_cause.split("\n\n"):
        if para.strip():
            body(doc, para.strip())

    # serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------- self-test
if __name__ == "__main__":
    demo_findings = [
        {
            "id": "F-01",
            "title": "Critical Isolation Failure Reported as Closed",
            "severity": "Critical",
            "observation": "The executive overview reports the cross-tenant IDOR as closed as of March. The penetration test shows the storage-layer equivalent remains open until Q4.",
            "assessment": "The application-layer authorization check is now the sole control preventing cross-tenant access. If it fails, the original exposure returns in full.",
            "traceability": [
                "Executive Overview 2.1 - reports critical risk remediated",
                "Pen Test PT-2025-007 - storage layer open, deferred to Q4",
            ],
            "position": "Do not report as closed until the storage layer is remediated or a formal risk acceptance is recorded.",
        },
        {
            "id": "F-02",
            "title": "Notification Decision Relies on Absent Logging",
            "severity": "Critical",
            "observation": "The no-notification decision rests on 'no evidence of exploitation.' Two documents confirm CloudTrail does not log the affected layer.",
            "traceability": [
                "Security Architecture 7.3 - CloudTrail excludes API Gateway",
                "Incident Response Plan Section 6 - confirms the gap",
            ],
            "position": "Re-open and re-document the decision on stated grounds, not on absent logs.",
        },
    ]

    data = build_report(
        title="ClearPath Financial Analytics Platform",
        subtitle="Cross-Document Findings, Root Cause, and Remediation Plan",
        prepared_by="Tyrell Mifflin - IAM / GRC Consulting",
        exec_summary="ClearPath operates a multi-tenant financial analytics platform. This assessment reviewed the full security documentation set to determine whether the stated posture is supported by the evidence.\n\nIt is not, in three material respects, each detailed below.",
        findings=demo_findings,
        root_cause="Six findings, one cause. ClearPath runs four independent assurance streams and nothing reconciles them.\n\nThe remedy is a reconciled findings register keyed on risk, not finding, with automated cross-stream checks.",
        bottom_line="The individual documents are competent. The failure is at the seams between them, where one artefact says closed and another says open.",
    )
    with open("/home/claude/demo_report.docx", "wb") as fh:
        fh.write(data)
    print(f"wrote demo_report.docx ({len(data)} bytes)")