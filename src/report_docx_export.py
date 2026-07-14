# src/report_docx_export.py
# Purpose: Convert a structured markdown risk assessment (as produced by
# vendor_doc_assessment.py) into a polished, styled Word document — headers,
# tables, and shaded "Position" callout boxes — matching a professional
# independent risk assessment report format.

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_heading(doc, text, level):
    """Add a styled heading, stripping markdown ** markers."""
    clean = text.replace("**", "").strip()
    p = doc.add_heading(level=level)
    run = p.add_run(clean)
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = GOLD
    return p


def _add_bold_aware_runs(paragraph, text, size=11, color=None, italic=False):
    """Add runs to an existing paragraph, rendering **bold** segments as bold."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(size)
        run.italic = italic
        if color:
            run.font.color.rgb = color


def _add_bold_aware_paragraph(doc, text, italic=False, size=11, color=None):
    """Add a new paragraph, rendering **bold** markdown segments as bold runs."""
    p = doc.add_paragraph()
    _add_bold_aware_runs(p, text, size=size, color=color, italic=italic)
    return p


def _is_position_callout_table(rows):
    """Detect a single-row pipe-table that is actually a Position callout,
    e.g. '| **Position:** ... |' followed by a '| --- |' separator, rather
    than a genuine multi-row data table."""
    content_rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r.strip())]
    if len(content_rows) != 1:
        return False
    text = content_rows[0].strip("|").strip()
    return bool(re.match(r"^\*\*Position", text, re.IGNORECASE)) or text.lower().startswith("position:")


def _add_callout(doc, text):
    """Add a shaded 'Position' callout box using a single-cell table."""
    clean = text.strip()
    if clean.startswith(">"):
        clean = clean[1:].strip()
    clean = clean.strip("|").strip()

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "F2E9CE")  # light gold tint
    cell.paragraphs[0].text = ""
    _add_bold_aware_runs(cell.paragraphs[0], clean, size=10.5, color=DARK)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()  # spacer after callout


def _add_markdown_table(doc, lines):
    """Render a markdown pipe table as a styled Word table with a navy header
    row — unless it is actually a Position callout box in disguise, in which
    case route it to the callout renderer instead."""
    rows = [l.strip() for l in lines if l.strip().startswith("|")]

    if _is_position_callout_table(rows):
        content_rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r.strip())]
        _add_callout(doc, content_rows[0])
        return

    # Drop the markdown separator row (---|---|---)
    rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r)]
    if not rows:
        return
    parsed = [
        [c.strip().replace("**", "") for c in r.strip("|").split("|")]
        for r in rows
    ]
    n_cols = max(len(row) for row in parsed)
    parsed = [row + [""] * (n_cols - len(row)) for row in parsed]

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(parsed[0]):
        hdr_cells[i].text = col_name
        _set_cell_shading(hdr_cells[i], "1B3A5C")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10)
    for row_data in parsed[1:]:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < n_cols:
                row_cells[i].text = val
                for p in row_cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
    doc.add_paragraph()


def markdown_report_to_docx(markdown_text: str, output_path: str, title: str = "Independent Risk Assessment"):
    """Convert a structured markdown risk report into a styled .docx file."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|"):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            _add_markdown_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:], level=1)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=2)
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=3)
        elif stripped.startswith(">"):
            _add_callout(doc, stripped)
        elif re.match(r"^-{3,}$", stripped):
            pass  # skip markdown horizontal rules (---) entirely, no visual value in Word
        elif stripped.startswith("====="):
            pass
        else:
            _add_bold_aware_paragraph(doc, stripped)

        i += 1

    if in_table and table_buffer:
        _add_markdown_table(doc, table_buffer)

    doc.save(output_path)
    return output_path